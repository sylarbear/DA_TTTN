from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
SRC = DOCS / "bao-cao-thuc-tap-tuan3.docx"
OUT = DOCS / "bao-cao-thuc-tap-tuan3_chuong2_viet_lai.docx"
DIAGRAMS = DOCS / "diagrams"


def paragraph_text(element):
    return "".join(t.text or "" for t in element.iter(qn("w:t"))).strip()


def remove_from_heading(doc, heading_text):
    body = doc._body._element
    children = list(body)
    start = None
    for idx, child in enumerate(children):
        if child.tag == qn("w:p") and paragraph_text(child) == heading_text:
            start = idx
            break
    if start is None:
        raise RuntimeError(f"Không tìm thấy heading: {heading_text}")
    for child in children[start:]:
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size=None, bold=None, italic=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_p(doc, text="", style="_doanVB", align=None, before=None, after=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    if before is not None:
        p.paragraph_format.space_before = Pt(before)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if first_line and style == "_doanVB":
        p.paragraph_format.first_line_indent = Inches(0.3)
    r = p.add_run(text)
    set_run_font(r, 13)
    return p


def add_heading(doc, text, style):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        r = p.add_run("• " + item)
        set_run_font(r, 13)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, 12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill="D9EAF7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(hdr[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], value, align=align)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()
    return table


def usable_width(doc):
    section = doc.sections[-1]
    return section.page_width - section.left_margin - section.right_margin


def add_figure(doc, filename, caption, width_inches=6.1):
    path = DIAGRAMS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, 12, bold=True, italic=True)
    cap.paragraph_format.space_after = Pt(6)


def add_chapter_2(doc):
    add_heading(doc, "LÝ THUYẾT ÁP DỤNG THEO LĨNH VỰC THỰC TẬP", "Heading 1")

    add_heading(doc, "Tổng quan về đề tài", "_level 1")
    add_p(doc, "Đề tài xây dựng website học tiếng Anh trực tuyến EngPath được thực hiện nhằm tạo ra một nền tảng hỗ trợ người học rèn luyện tiếng Anh theo lộ trình rõ ràng. Hệ thống tập trung vào các hoạt động học tập thường gặp gồm học từ vựng theo chủ đề, đọc bài học, làm bài kiểm tra, ôn tập flashcard, luyện nói với phản hồi tự động và theo dõi tiến độ cá nhân.")
    add_p(doc, "Đề tài phù hợp với lĩnh vực thực tập Công nghệ thông tin vì yêu cầu vận dụng đồng thời nhiều nhóm kiến thức: lập trình web phía máy chủ, thiết kế giao diện người dùng, thiết kế cơ sở dữ liệu quan hệ, phân quyền tài khoản, bảo mật form, xử lý AJAX, tích hợp API bên ngoài và kiểm thử chức năng. Sản phẩm được triển khai theo mô hình web MVC thuần PHP, chạy được trên môi trường XAMPP với Apache và MySQL.")
    add_p(doc, "Hệ thống hướng đến hai nhóm người dùng chính. Người học sử dụng các chức năng học tập, luyện tập, nâng cấp Pro và gửi hỗ trợ. Quản trị viên sử dụng khu vực admin riêng để theo dõi dashboard, quản lý user, khóa học, câu hỏi, đơn nâng cấp, ticket hỗ trợ và cấu hình API.")

    add_table(
        doc,
        ["Thành phần", "Công nghệ áp dụng", "Vai trò trong hệ thống"],
        [
            ["Backend", "PHP 8.x, mô hình MVC", "Xử lý request, điều hướng controller, gọi model và render view."],
            ["Database", "MySQL, PDO, utf8mb4", "Lưu người dùng, nội dung học, kết quả test, speaking, membership và ticket."],
            ["Frontend", "HTML5, CSS3, JavaScript, AJAX", "Xây dựng giao diện, tương tác form, gọi API nội bộ và xử lý trải nghiệm người dùng."],
            ["Xác thực", "Session PHP, password_hash, Google OAuth 2.0", "Đăng nhập tài khoản thường, đăng nhập Google và duy trì phiên làm việc."],
            ["AI/Speaking", "Web Speech API, OpenAI GPT API", "Nhận dạng giọng nói, chấm điểm speaking và chatbot hỗ trợ học tập."],
            ["Quản trị", "AdminController, Middleware, CSRF", "Tách riêng quyền admin và bảo vệ các thao tác quản trị."],
        ],
        widths=[1.35, 2.0, 3.0],
    )

    add_heading(doc, "Kiến trúc tổng quan của hệ thống", "_level 1")
    add_p(doc, "EngPath được tổ chức theo kiến trúc MVC. Toàn bộ request đi qua file public/index.php, sau đó lớp Router xác định controller và action tương ứng. Controller tiếp nhận dữ liệu đầu vào, kiểm tra quyền truy cập bằng Middleware, gọi Model để truy vấn cơ sở dữ liệu và trả kết quả về View hoặc JSON.")
    add_p(doc, "Tầng core của hệ thống bao gồm các lớp App, Router, Controller, Model, Middleware, CSRF, Env, Logger, RateLimiter, GoogleOAuth, OpenAIService và StreakService. Việc tách các lớp lõi giúp mã nguồn có cấu trúc rõ ràng, dễ bảo trì và hạn chế lặp lại logic giữa nhiều module.")
    add_figure(doc, "architecture.drawio.png", "Hình 2.1. Kiến trúc tổng quan hệ thống EngPath", 6.3)

    add_heading(doc, "Cơ sở lý thuyết", "_level 1")
    add_heading(doc, "Ngôn ngữ lập trình PHP", "_level2")
    add_p(doc, "PHP là ngôn ngữ lập trình phía máy chủ được sử dụng phổ biến trong phát triển website động. Trong đề tài, PHP đảm nhiệm việc xử lý đăng ký, đăng nhập, phân quyền, truy vấn cơ sở dữ liệu, xử lý biểu mẫu, tạo phản hồi JSON và điều hướng giao diện. Việc sử dụng PHP thuần giúp sinh viên hiểu rõ luồng xử lý request-response thay vì phụ thuộc quá nhiều vào framework có sẵn.")
    add_bullets(doc, [
        "Sử dụng cú pháp đơn giản, dễ triển khai trên XAMPP.",
        "Kết hợp tốt với MySQL thông qua PDO và prepared statement.",
        "Hỗ trợ session để duy trì trạng thái đăng nhập.",
        "Có thể tổ chức mã nguồn theo mô hình MVC để tăng tính bảo trì.",
    ])

    add_heading(doc, "Mô hình MVC", "_level2")
    add_p(doc, "MVC là mô hình chia ứng dụng thành ba thành phần chính: Model, View và Controller. Model chịu trách nhiệm thao tác dữ liệu; View chịu trách nhiệm hiển thị giao diện; Controller tiếp nhận yêu cầu từ người dùng và điều phối xử lý. Trong EngPath, các controller như AuthController, TopicController, TestController, SpeakingController, MembershipController và AdminController giúp tách riêng từng nhóm nghiệp vụ.")
    add_p(doc, "Mô hình này giúp quá trình mở rộng chức năng thuận lợi hơn. Ví dụ khi bổ sung thêm trang quản trị hoặc chức năng speaking, lập trình viên có thể tạo controller, view và model tương ứng mà không làm xáo trộn toàn bộ hệ thống.")

    add_heading(doc, "Hệ quản trị cơ sở dữ liệu MySQL", "_level2")
    add_p(doc, "MySQL được sử dụng để lưu trữ dữ liệu chính của hệ thống. Cơ sở dữ liệu của đề tài dùng charset utf8mb4 nhằm hỗ trợ tiếng Việt và ký tự đặc biệt. Các truy vấn trong code được thực hiện thông qua PDO và prepared statement để giảm rủi ro SQL Injection.")
    add_p(doc, "Database hiện tại gồm 23 bảng, phục vụ các nhóm nghiệp vụ: tài khoản người dùng, nội dung học, bài kiểm tra, luyện nói, ngữ pháp, tiến độ học tập, membership, mã kích hoạt, ticket hỗ trợ và lịch sử điểm kinh nghiệm.")

    add_heading(doc, "HTML5, CSS3 và JavaScript", "_level2")
    add_p(doc, "HTML5 được dùng để xây dựng cấu trúc giao diện; CSS3 dùng để định dạng, tạo layout responsive, card, bảng, nút và các trạng thái tương tác; JavaScript dùng để xử lý hành vi phía trình duyệt như mở modal, gửi form AJAX, ghi âm speaking, cập nhật biểu đồ và hiển thị phản hồi động.")
    add_p(doc, "Giao diện hiện tại của EngPath được thiết kế lại theo hướng sáng, rõ ràng và thân thiện hơn với người học. Các file CSS chính gồm style.css, components.css và pages.css; các file JavaScript chính gồm app.js, dashboard.js và speaking.js.")

    add_heading(doc, "Xác thực và phân quyền", "_level2")
    add_p(doc, "Hệ thống hỗ trợ đăng ký, đăng nhập bằng username/password và đăng nhập bằng Google OAuth 2.0. Mật khẩu được xử lý bằng password_hash và password_verify. Sau khi đăng nhập thành công, hệ thống lưu thông tin cần thiết vào session để nhận diện người dùng, vai trò và trạng thái membership.")
    add_p(doc, "Middleware được sử dụng để kiểm tra quyền truy cập. Người dùng chưa đăng nhập sẽ bị chuyển về trang đăng nhập khi truy cập các chức năng yêu cầu tài khoản. Admin chỉ được sử dụng khu vực quản trị; các trang học tập thông thường dành cho học viên.")

    add_heading(doc, "Web Speech API và OpenAI GPT API", "_level2")
    add_p(doc, "Web Speech API cho phép trình duyệt nhận dạng giọng nói và chuyển lời nói của người học thành văn bản. Trong chức năng luyện nói, JavaScript ghi nhận transcript và confidence, sau đó gửi dữ liệu về SpeakingController để chấm điểm.")
    add_p(doc, "OpenAI GPT API được tích hợp thông qua lớp OpenAIService. Khi API khả dụng, hệ thống gửi transcript và câu trả lời mẫu để AI đánh giá Accuracy, Fluency, Pronunciation, Overall Score và Feedback. Nếu chưa cấu hình API key, hệ thống vẫn có cơ chế chấm điểm cục bộ dựa trên so khớp từ khóa và độ tương đồng văn bản.")

    add_heading(doc, "Thanh toán QR và xử lý đơn nâng cấp Pro", "_level2")
    add_p(doc, "Chức năng nâng cấp Pro trong phiên bản hiện tại sử dụng hình thức chuyển khoản ngân hàng qua QR. Người dùng chọn gói Pro, nhập nội dung chuyển khoản và tạo đơn pending. Sau khi kiểm tra giao dịch thực tế, admin duyệt đơn trong trang quản trị để cập nhật trạng thái completed và nâng membership của user lên Pro.")
    add_p(doc, "Cách triển khai này phù hợp với phạm vi đồ án vì dễ kiểm soát dữ liệu, không phụ thuộc vào webhook thanh toán thật và vẫn thể hiện được đầy đủ quy trình nghiệp vụ: tạo đơn, chờ duyệt, duyệt/từ chối, tính ngày hết hạn gói Pro.")

    add_heading(doc, "Gamification trong học tập", "_level2")
    add_p(doc, "Hệ thống sử dụng XP, level, streak, daily goal và leaderboard để tạo động lực học tập. Khi người học hoàn thành các hoạt động như học từ vựng, làm bài test, luyện nói hoặc hoàn thành bài học, hệ thống cập nhật tiến độ và cộng điểm kinh nghiệm tương ứng. Dashboard cá nhân giúp người học theo dõi quá trình học của mình theo thời gian.")


def add_chapter_3(doc):
    doc.add_page_break()
    add_heading(doc, "NỘI DUNG THỰC TẬP", "Heading 1")

    add_heading(doc, "Phân tích yêu cầu hệ thống", "_level 1")
    add_p(doc, "Từ mục tiêu xây dựng website học tiếng Anh trực tuyến, hệ thống EngPath được phân tích thành hai nhóm chức năng lớn: nhóm chức năng dành cho người học và nhóm chức năng dành cho quản trị viên. Người học cần thao tác thuận tiện, giao diện dễ hiểu và phản hồi nhanh; quản trị viên cần có khu vực riêng để quản lý dữ liệu và theo dõi hoạt động.")
    add_heading(doc, "Yêu cầu chức năng", "_level2")
    add_bullets(doc, [
        "Người dùng có thể đăng ký, đăng nhập, đăng xuất, đăng nhập bằng Google và quản lý hồ sơ cá nhân.",
        "Người học có thể xem danh sách khóa học/chủ đề, học từ vựng, đọc bài học, ôn flashcard và lưu bookmark.",
        "Người học có thể làm bài test Quiz, Listening, Reading và xem kết quả sau khi nộp bài.",
        "Người học có thể luyện nói bằng microphone, nhận điểm chấm và feedback.",
        "Người học có thể theo dõi dashboard cá nhân, XP, streak, level và bảng xếp hạng.",
        "Người học có thể chọn gói Pro, tạo đơn chuyển khoản QR và gửi ticket hỗ trợ.",
        "Admin có thể quản lý user, khóa học/chủ đề, câu hỏi, đơn nâng cấp, ticket và cấu hình OpenAI API key.",
    ])
    add_heading(doc, "Yêu cầu phi chức năng", "_level2")
    add_bullets(doc, [
        "Giao diện phải rõ ràng, dễ dùng, responsive trên desktop và thiết bị di động.",
        "Dữ liệu tiếng Việt phải hiển thị đúng mã hóa UTF-8/utf8mb4.",
        "Các thao tác quan trọng cần kiểm tra quyền truy cập và chống gửi request sai phương thức.",
        "Các truy vấn cơ sở dữ liệu cần dùng prepared statement để tăng tính an toàn.",
        "Hệ thống cần chạy ổn định trên XAMPP, Apache và MySQL trong môi trường local.",
    ])

    add_heading(doc, "Sơ đồ Use Case", "_level 1")
    add_p(doc, "Sơ đồ Use Case tổng quan thể hiện các tác nhân chính và nhóm chức năng của hệ thống. Actor Learner đại diện cho người học sử dụng các chức năng học tập, luyện nói, nâng cấp Pro và hỗ trợ. Actor Admin đại diện cho quản trị viên sử dụng khu vực quản trị. Google OAuth và OpenAI API là các dịch vụ ngoài được hệ thống tích hợp.")
    add_figure(doc, "usecase_overview.png", "Hình 3.1. Sơ đồ Use Case tổng quan hệ thống EngPath", 5.9)

    add_heading(doc, "Sơ đồ Sequence", "_level 1")
    add_heading(doc, "Luồng đăng nhập", "_level2")
    add_p(doc, "Luồng đăng nhập bắt đầu khi người dùng gửi username và password từ trình duyệt. AuthController gọi User Model để tìm tài khoản, kiểm tra mật khẩu bằng password_verify, tạo lại session ID để hạn chế session fixation, lưu thông tin người dùng vào session và điều hướng về trang phù hợp với vai trò.")
    add_figure(doc, "sequence_login.drawio.png", "Hình 3.2. Sequence Diagram - Đăng nhập", 6.1)

    add_heading(doc, "Luồng luyện nói AI", "_level2")
    add_p(doc, "Ở chức năng luyện nói, trình duyệt sử dụng Web Speech API để nhận dạng giọng nói. Transcript được gửi về SpeakingController qua endpoint /speaking/score. Controller lấy prompt tương ứng, gọi OpenAIService nếu có API key, lưu kết quả vào speaking_attempts, cập nhật tiến độ user_progress và cộng XP khi điểm tổng đạt điều kiện.")
    add_figure(doc, "sequence_speaking.drawio.png", "Hình 3.3. Sequence Diagram - Luyện nói AI", 6.1)

    add_heading(doc, "Luồng nâng cấp Pro bằng QR", "_level2")
    add_p(doc, "Luồng nâng cấp Pro gồm hai giai đoạn. Giai đoạn đầu, người dùng chọn gói và tạo đơn chuyển khoản, hệ thống lưu đơn với trạng thái pending. Giai đoạn sau, admin kiểm tra giao dịch ngân hàng và duyệt đơn. Khi duyệt thành công, hệ thống cập nhật membership_orders sang completed và cập nhật users.membership thành pro kèm ngày hết hạn.")
    add_figure(doc, "sequence_purchase.drawio.png", "Hình 3.4. Sequence Diagram - Nâng cấp Pro bằng QR và Admin duyệt", 6.1)

    add_heading(doc, "Sơ đồ State", "_level 1")
    add_p(doc, "Sơ đồ trạng thái mô tả vòng đời của membership, đơn nâng cấp Pro và ticket hỗ trợ. Membership chuyển từ Free sang Pro khi admin duyệt đơn QR và chuyển về Free khi hết hạn hoặc admin hạ cấp. Đơn nâng cấp đi từ Pending sang Completed hoặc Cancelled. Ticket hỗ trợ đi từ Open sang In Progress, Resolved và Closed.")
    add_figure(doc, "state_diagrams.drawio.png", "Hình 3.5. State Diagram - Membership, đơn nâng cấp và ticket", 6.1)

    add_heading(doc, "Thiết kế cơ sở dữ liệu", "_level 1")
    add_p(doc, "Cơ sở dữ liệu english_master được thiết kế theo mô hình quan hệ với 23 bảng. Các bảng được chia theo nhóm nghiệp vụ để dễ quản lý: người dùng, nội dung học, kiểm tra, luyện nói, membership, hỗ trợ và gamification. Các bảng có quan hệ khóa ngoại giúp đảm bảo tính toàn vẹn dữ liệu khi xóa hoặc cập nhật bản ghi.")
    add_figure(doc, "erd.drawio.png", "Hình 3.6. Sơ đồ cơ sở dữ liệu EngPath - 23 bảng", 6.3)
    add_table(
        doc,
        ["Nhóm dữ liệu", "Các bảng chính", "Mục đích"],
        [
            ["Người dùng", "users, user_progress, xp_history, bookmarks", "Quản lý tài khoản, tiến độ, điểm XP và từ đã lưu."],
            ["Nội dung học", "topics, vocabularies, lessons, lesson_contents, lesson_reviews", "Lưu chủ đề, từ vựng, bài học, nội dung bài học và đánh giá."],
            ["Kiểm tra", "tests, questions, test_results, user_answers", "Tạo bài test, câu hỏi, kết quả làm bài và đáp án người dùng."],
            ["Ngữ pháp", "grammar_lessons, grammar_questions", "Quản lý bài ngữ pháp và câu hỏi luyện tập."],
            ["Luyện nói", "speaking_prompts, speaking_attempts", "Lưu prompt, câu trả lời mẫu, transcript, điểm và feedback speaking."],
            ["Pro và giao dịch", "membership_plans, membership_orders, activation_codes, wallet_transactions", "Quản lý gói Pro, đơn nâng cấp, mã kích hoạt và bảng giao dịch kế thừa."],
            ["Hỗ trợ", "support_tickets", "Lưu ticket hỗ trợ, yêu cầu hủy đơn và phản hồi admin."],
        ],
        widths=[1.35, 2.6, 2.65],
    )

    add_heading(doc, "Xây dựng giao diện ứng dụng", "_level 1")
    add_p(doc, "Giao diện người học được thiết kế theo hướng hiện đại, sáng, có nhiều khoảng trắng và điều hướng rõ ràng. Header của người học gồm các mục Trang chủ, Khóa học, Bài test, Luyện nói, Pro và menu mở rộng cho Dashboard, Ngữ pháp, Xếp hạng, Từ đã lưu, Tìm kiếm và Hỗ trợ.")
    add_bullets(doc, [
        "Trang chủ giới thiệu lộ trình học, các tính năng chính và lời kêu gọi bắt đầu học miễn phí.",
        "Trang khóa học/chủ đề hiển thị danh sách topic theo cấp độ, số lượng từ vựng, bài học và bài test.",
        "Trang chi tiết chủ đề gom từ vựng, bài học, test và lối vào flashcard.",
        "Trang bài test phân loại Quiz, Listening và Reading, hỗ trợ nộp bài bằng AJAX và xem kết quả.",
        "Trang luyện nói cung cấp prompt, ghi âm, transcript, điểm số và feedback.",
        "Trang Pro hiển thị bảng giá, quyền lợi Free/Pro, QR chuyển khoản và lịch sử đơn nâng cấp.",
        "Trang Dashboard, Leaderboard và Profile giúp người học theo dõi tiến độ và thông tin cá nhân.",
    ])

    add_heading(doc, "Giao diện quản trị", "_level 1")
    add_p(doc, "Giao diện admin được tách riêng khỏi trải nghiệm học viên. Khi tài khoản có role admin đăng nhập, header và khu vực điều hướng chuyển sang các chức năng quản trị. Việc tách này giúp admin tập trung quản lý hệ thống, không bị lẫn với các màn hình học tập của user.")
    add_bullets(doc, [
        "Dashboard Admin thống kê học viên, Pro users, chủ đề, bài test, câu hỏi, lượt làm bài, đơn pending và ticket chờ xử lý.",
        "Quản lý Users cho phép tìm kiếm, cập nhật thông tin, role, membership và xóa user không phải admin.",
        "Quản lý Khóa học/Chủ đề cho phép thêm, sửa chủ đề và theo dõi số bài học, từ vựng, bài test.",
        "Quản lý Câu hỏi cho phép xem test, thêm/sửa/xóa câu hỏi và đáp án.",
        "Quản lý Đơn nâng cấp cho phép duyệt hoặc từ chối đơn Pro đang chờ.",
        "Quản lý Tickets cho phép phản hồi, đổi trạng thái và xử lý yêu cầu hủy đơn.",
        "Cài đặt cho phép cấu hình OpenAI API key dùng cho Speaking AI và Chatbot.",
    ])

    add_heading(doc, "Xây dựng chức năng", "_level 1")
    add_heading(doc, "Quản lý tài khoản và phân quyền", "_level2")
    add_p(doc, "AuthController xử lý đăng ký, đăng nhập, đăng xuất và đăng nhập Google. User Model chịu trách nhiệm tạo tài khoản, xác thực mật khẩu, kiểm tra username/email và tìm hoặc tạo tài khoản từ Google OAuth. Middleware kiểm tra trạng thái đăng nhập, quyền admin và trạng thái Pro.")

    add_heading(doc, "Quản lý nội dung học", "_level2")
    add_p(doc, "TopicController, LessonController, GrammarController, TestController và Vocabulary Model phục vụ phần nội dung học tập. Người học có thể xem chủ đề, học từ vựng, xem bài học, làm quiz ngữ pháp, làm bài test và ôn tập bằng flashcard. Các hoạt động học tập được cập nhật vào user_progress để phục vụ dashboard.")

    add_heading(doc, "Luyện nói và phản hồi AI", "_level2")
    add_p(doc, "SpeakingController cung cấp danh sách prompt, trang luyện nói và endpoint chấm điểm. public/js/speaking.js điều khiển microphone, nhận dạng giọng nói, gửi transcript và hiển thị kết quả. SpeakingAttempt Model lưu điểm Accuracy, Fluency, Pronunciation, Overall Score và feedback vào bảng speaking_attempts.")

    add_heading(doc, "Nâng cấp Pro và hỗ trợ", "_level2")
    add_p(doc, "MembershipController hiển thị danh sách gói Pro, lịch sử đơn và xử lý tạo đơn chuyển khoản. AdminController xử lý duyệt/từ chối đơn. SupportController cho phép người dùng gửi ticket hỗ trợ, gửi yêu cầu hủy đơn pending và kiểm tra điều kiện hủy. Admin có thể phản hồi ticket, đổi trạng thái và duyệt hủy đơn từ khu vực quản trị.")

    add_heading(doc, "Gamification và dashboard", "_level2")
    add_p(doc, "StreakService quản lý streak, XP, daily goal và lịch sử hoạt động. DashboardController lấy dữ liệu tiến độ học tập để hiển thị cho người học. LeaderboardController hiển thị bảng xếp hạng dựa trên tổng XP, giúp tăng tính cạnh tranh và động lực học tập.")

    add_heading(doc, "Kiểm thử", "_level 1")
    add_p(doc, "Quá trình kiểm thử được thực hiện trên môi trường XAMPP với Apache và MySQL. Các chức năng chính được kiểm tra theo từng nhóm: xác thực tài khoản, học tập, bài test, speaking, membership, ticket và admin.")
    add_bullets(doc, [
        "Kiểm tra đăng ký, đăng nhập, đăng xuất và phân quyền admin/student.",
        "Kiểm tra hiển thị tiếng Việt, nội dung chủ đề, từ vựng, bài học và flashcard.",
        "Kiểm tra làm bài test, lưu kết quả và xem chi tiết đáp án.",
        "Kiểm tra luyện nói: cấp quyền microphone, nhận transcript, chấm điểm và lưu attempt.",
        "Kiểm tra tạo đơn Pro, trạng thái pending, admin duyệt/từ chối và cập nhật ngày hết hạn.",
        "Kiểm tra ticket hỗ trợ, phản hồi admin và thay đổi trạng thái.",
        "Kiểm tra giao diện responsive và các thao tác AJAX không làm vỡ bố cục.",
    ])


def add_chapter_4(doc):
    doc.add_page_break()
    add_heading(doc, "KẾT LUẬN VÀ KIẾN NGHỊ", "Heading 1")

    add_heading(doc, "Kết quả đạt được", "_level 1")
    add_p(doc, "Sau quá trình thực hiện, đề tài đã xây dựng được website học tiếng Anh trực tuyến EngPath với đầy đủ các nhóm chức năng cốt lõi cho người học và quản trị viên. Hệ thống có thể chạy trên môi trường local, sử dụng cơ sở dữ liệu MySQL đầy đủ, giao diện đã được làm mới và luồng nghiệp vụ chính hoạt động theo đúng mục tiêu đề ra.")
    add_heading(doc, "Ưu điểm", "_level2")
    add_bullets(doc, [
        "Tổ chức mã nguồn theo mô hình MVC, tách rõ controller, model, view và core.",
        "Hoàn thiện các chức năng học tập: chủ đề, từ vựng, bài học, flashcard, bài test, ngữ pháp và bookmark.",
        "Tích hợp chức năng luyện nói sử dụng Web Speech API và OpenAI GPT API.",
        "Có dashboard học tập, XP, streak, level và bảng xếp hạng để tăng động lực học.",
        "Có quy trình nâng cấp Pro bằng QR và admin duyệt thủ công, phù hợp phạm vi đồ án.",
        "Có khu vực admin riêng để quản lý người dùng, chủ đề, câu hỏi, đơn nâng cấp, ticket và API key.",
        "Giao diện người dùng được cải thiện theo hướng hiện đại, dễ thao tác và responsive.",
    ])

    add_heading(doc, "Hạn chế", "_level 1")
    add_bullets(doc, [
        "Thanh toán Pro chưa tích hợp cổng thanh toán tự động hoặc webhook xác nhận giao dịch thực tế.",
        "Nội dung học tập vẫn cần bổ sung thêm chủ đề, bài học, câu hỏi listening/reading và dữ liệu speaking.",
        "Chức năng speaking phụ thuộc vào trình duyệt hỗ trợ Web Speech API và quyền truy cập microphone.",
        "Hệ thống mới triển khai ở môi trường local/XAMPP, chưa tối ưu đầy đủ cho môi trường hosting thật.",
        "Chưa có ứng dụng mobile native, người dùng chủ yếu truy cập qua trình duyệt web.",
    ])

    add_heading(doc, "Kiến nghị và hướng phát triển", "_level 1")
    add_bullets(doc, [
        "Triển khai hệ thống lên hosting/VPS, cấu hình HTTPS, backup database và logging vận hành.",
        "Tích hợp cổng thanh toán tự động như VNPay, Momo hoặc webhook ngân hàng để xác nhận giao dịch real-time.",
        "Bổ sung thêm nội dung học tập theo nhiều cấp độ, chủ đề và kỹ năng.",
        "Hoàn thiện thuật toán ôn tập từ vựng theo Spaced Repetition System để cá nhân hóa lịch học.",
        "Mở rộng AI chatbot thành trợ lý hội thoại tiếng Anh có ngữ cảnh theo bài học.",
        "Phát triển phiên bản mobile hoặc PWA để người học sử dụng thuận tiện hơn.",
        "Bổ sung báo cáo thống kê nâng cao cho admin: doanh thu, tỉ lệ hoàn thành bài học, tỉ lệ chuyển đổi Pro.",
    ])

    doc.add_page_break()
    add_heading(doc, "TÀI LIỆU THAM KHẢO", "Heading 1")
    add_p(doc, "Tài liệu tham khảo được sử dụng trong quá trình tìm hiểu công nghệ và xây dựng hệ thống:", style="_doanVB")
    refs = [
        "PHP Manual. Địa chỉ: https://www.php.net/docs.php",
        "MySQL 8.0 Reference Manual. Địa chỉ: https://dev.mysql.com/doc/refman/8.0/en/",
        "MDN Web Docs - HTML, CSS, JavaScript. Địa chỉ: https://developer.mozilla.org/",
        "MDN Web Docs - Web Speech API. Địa chỉ: https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API",
        "Google Identity - OAuth 2.0 Documentation. Địa chỉ: https://developers.google.com/identity/protocols/oauth2",
        "OpenAI API Documentation. Địa chỉ: https://platform.openai.com/docs/",
        "Chart.js Documentation. Địa chỉ: https://www.chartjs.org/docs/",
        "Font Awesome Documentation. Địa chỉ: https://fontawesome.com/docs/",
    ]
    for idx, ref in enumerate(refs, 1):
        add_p(doc, f"[{idx}] {ref}", style="Normal", first_line=False)


def main():
    doc = Document(SRC)
    remove_from_heading(doc, "LÝ THUYẾT ÁP DỤNG THEO LĨNH VỰC THỰC TẬP")
    doc.add_page_break()
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
