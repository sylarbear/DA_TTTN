# -*- coding: utf-8 -*-
"""
Generate bao_cao_thuc_tap.docx using the EXACT template styles and formatting.
Template: Mau-bao-cao-chuyen-de-thuc-tap_New.docx
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

DOCS_DIR = r"g:\xamp\htdocs\DATN\docs"
TEMPLATE = os.path.join(DOCS_DIR, "Mau-bao-cao-chuyen-de-thuc-tap_New.docx")
SCREENSHOTS_DIR = os.path.join(DOCS_DIR, "screenshots")
UML_DIR = os.path.join(DOCS_DIR, "uml_images")
OUTPUT = os.path.join(DOCS_DIR, "bao_cao_thuc_tap_v2.docx")

# Load template
doc = Document(TEMPLATE)

# Clear all content but keep styles + sections
body = doc.element.body
sect_elements = []
for child in list(body):
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'sectPr':
        sect_elements.append(child)
    else:
        body.remove(child)

# Re-append section properties
for sect in sect_elements:
    body.append(sect)

# ============================================================
# HELPERS using TEMPLATE styles exactly
# ============================================================
def heading1(text):
    """Chapter heading - Heading 1 style (20pt bold center caps)"""
    p = doc.add_paragraph(text, 'Heading 1')
    return p

def level1(text):
    """Section heading - _level 1 (14pt, indent 0.3in)"""
    return doc.add_paragraph(text, '_level 1')

def level2(text):
    """Sub-section - _level2 (14pt, indent 0.39in)"""
    return doc.add_paragraph(text, '_level2')

def level3(text):
    """Sub-sub-section - _level3"""
    return doc.add_paragraph(text, '_level3')

def body_text(text):
    """Body paragraph - _doanVB (13pt, justify, first indent, 1.2 spacing)"""
    return doc.add_paragraph(text, '_doanVB')

def bullet(text):
    """List item - List Paragraph"""
    p = doc.add_paragraph(text, 'List Paragraph')
    return p

def normal(text, bold=False, size=None, align=None):
    """Normal paragraph"""
    p = doc.add_paragraph(style='Normal')
    run = p.add_run(text)
    if bold:
        run.bold = True
    if size:
        run.font.size = Pt(size)
    if align is not None:
        p.paragraph_format.alignment = align
    return p

def normal_center(text, bold=False, size=None):
    return normal(text, bold, size, WD_ALIGN_PARAGRAPH.CENTER)

def normal_right(text, bold=False, size=None):
    return normal(text, bold, size, WD_ALIGN_PARAGRAPH.RIGHT)

def empty():
    doc.add_paragraph('', 'Normal')

def page_break():
    doc.add_page_break()

def add_image(img_path, caption=None, width=Inches(5.5)):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = normal_center(caption, size=11)
            p.runs[0].italic = True
    else:
        normal(f"[Hình: {os.path.basename(img_path)}]")

# ============================================================
# TRANG BÌA (match template exactly)
# ============================================================
normal_center("TRƯỜNG CAO ĐẲNG CÔNG THƯƠNG TPHCM", size=18)
p = normal_center("KHOA CÔNG NGHỆ THÔNG TIN", size=18)
p.runs[0].bold = True

empty()
empty()
empty()

normal_center("BÁO CÁO THỰC TẬP", True, 28)
normal_center("TỐT NGHIỆP", True, 28)

empty()

normal_center("ĐỀ TÀI: XÂY DỰNG WEBSITE HỌC TIẾNG ANH TRỰC TUYẾN", True, 16)
normal_center("ENGLISH LEARNING", True, 16)

empty()
empty()

normal("GV hướng dẫn: Vũ Thị Hường", size=18)
normal("Sinh viên thực hiện: Phan Quang Thuật", size=18)
normal("Mã số sinh viên: 2120110351", size=18)
normal("Lớp: CCQ2011E", size=18)
normal("Khoá: K44", size=18)

empty()
empty()
empty()
empty()

normal_center("TPHCM, tháng 4 năm 2026", True, 15)
page_break()

# ============================================================
# LỜI CẢM ƠN
# ============================================================
heading1("LỜI CẢM ƠN")

body_text("Trong quá trình thực hiện đề tài \"Xây dựng website học tiếng Anh trực tuyến English Learning\", em đã nhận được rất nhiều sự giúp đỡ và hướng dẫn tận tình.")
body_text("Trước hết, em xin gửi lời cảm ơn chân thành đến cô Vũ Thị Hường – Giảng viên hướng dẫn, người đã tận tình chỉ bảo, định hướng và đóng góp những ý kiến quý báu trong suốt quá trình em thực hiện đề tài. Những góp ý của cô đã giúp em hoàn thiện sản phẩm một cách tốt nhất.")
body_text("Em xin cảm ơn quý Thầy Cô trong Khoa Công nghệ Thông tin – Trường Cao đẳng Công Thương TPHCM đã truyền đạt cho em những kiến thức nền tảng vững chắc trong suốt thời gian học tập tại trường, tạo điều kiện cho em có thể áp dụng vào thực tế.")
body_text("Cuối cùng, em xin gửi lời cảm ơn đến gia đình, bạn bè đã luôn động viên, hỗ trợ em trong suốt quá trình học tập và thực hiện đề tài.")
body_text("Trong quá trình thực hiện, do kiến thức và kinh nghiệm còn hạn chế nên đề tài không tránh khỏi những thiếu sót. Em rất mong nhận được sự góp ý của quý Thầy Cô để em có thể hoàn thiện hơn.")

empty()
normal_right("TPHCM, tháng 4 năm 2026")
normal_right("Sinh viên thực hiện", True)
empty()
normal_right("Phan Quang Thuật", True)
page_break()

# ============================================================
# NHẬN XÉT GVHD (để trống cho cô điền)
# ============================================================
normal_center("NHẬN XÉT KẾT QUẢ THỰC TẬP", True, 16)
p = normal_center("(Dành cho Giảng viên hướng dẫn sinh viên thực tập)")
p.runs[0].italic = True

empty()
bullet("Họ tên sinh viên thực tập: Phan Quang Thuật")
bullet("Mã số sinh viên: 2120110351")
bullet("Lớp: CCQ2011E – Khóa: K44")
bullet("Giảng viên hướng dẫn thực tập: Vũ Thị Hường")
bullet("Sau thời gian sinh viên Phan Quang Thuật thực tập tốt nghiệp, tôi có nhận xét sau:")
empty()

normal("Về ý thức, thái độ của sinh viên:", True)
for _ in range(3):
    normal(".......................................................................................................................................")

normal("Về năng lực chuyên môn:", True)
for _ in range(3):
    normal(".......................................................................................................................................")

normal("Kết luận:", True)
normal(".......................................................................................................................................")
normal("Nhận xét:", True)
normal(".......................................................................................................................................")
normal("Điểm:", True)
normal(".......................................................................................................................................")

empty()
normal_right("TPHCM, Ngày…… tháng…… năm 2026")
normal_right("Giảng viên hướng dẫn", True)
empty()
empty()
normal_right("Vũ Thị Hường", True)
page_break()

# ============================================================
# ĐỀ CƯƠNG THỰC TẬP
# ============================================================
heading1("ĐỀ CƯƠNG THỰC TẬP")

normal("1. Nhiệm vụ / Đề tài:", True)
body_text("Xây dựng website học tiếng Anh trực tuyến English Learning.")

normal("2. Mục tiêu công việc / đề tài:", True)
bullet("Đề tài thuộc lĩnh vực phát triển ứng dụng web, giải quyết bài toán hỗ trợ người dùng học tiếng Anh trực tuyến thông qua các phương pháp học tương tác: từ vựng, ngữ pháp, luyện nghe, đọc, nói.")
bullet("Ứng dụng mang ý nghĩa thiết thực trong giáo dục, giúp người học tiếp cận việc học tiếng Anh một cách dễ dàng, miễn phí hoặc với chi phí thấp, mọi lúc mọi nơi thông qua trình duyệt web.")
bullet("Nắm vững kiến thức về xây dựng website theo mô hình MVC sử dụng PHP thuần, MySQL, HTML/CSS/JavaScript.")
bullet("Xây dựng một trang web thân thiện, tích hợp AI đánh giá kỹ năng nói cho người học tiếng Anh.")
bullet("Xây dựng hệ thống quản lý membership (ví điện tử, nạp tiền, mua gói) chuyên nghiệp và bảo mật.")

normal("3. Đối tượng nghiên cứu:", True)
bullet("Đối tượng nghiên cứu của đề tài là lập trình web nói chung và ngôn ngữ lập trình PHP, HTML, CSS, JavaScript nói riêng.")
bullet("Cơ sở dữ liệu MySQL, quản lý phiên bản với Git/GitHub.")
bullet("Tích hợp API: Google OAuth 2.0, OpenAI GPT, Web Speech API, VietQR.")
bullet("Nghiên cứu nhu cầu của người dùng học tiếng Anh online để xây dựng website thân thiện, dễ sử dụng.")

normal("4. Nội dung chính của thực tập:", True)
normal("Chương 1: Tổng quan về đề tài")
normal("Chương 2: Lý thuyết áp dụng theo lĩnh vực thực tập")
bullet("Nghiên cứu tổng quan về đề tài.")
bullet("Nghiên cứu cơ sở lý thuyết của đề tài.")
normal("Chương 3: Nội dung thực tập")
bullet("Tiến hành phân tích và thiết kế.")
bullet("Phân tích thiết kế hệ thống.")
bullet("Xây dựng cơ sở dữ liệu.")
bullet("Xây dựng giao diện ứng dụng trên nền web.")
bullet("Xây dựng chức năng cho ứng dụng.")
bullet("Kiểm thử và kiểm tra, sửa lỗi.")
normal("Chương 4: Kết luận và kiến nghị")
bullet("Kết quả thực hiện: Ưu điểm, Hạn chế.")
bullet("Hướng phát triển trong tương lai.")

normal("5. Tiến độ thực hiện thực tập:", True)
table = doc.add_table(rows=6, cols=4)
table.style = 'Table Grid'
headers = ['TT', 'Thời gian', 'Nội dung công việc', 'Kết quả dự kiến đạt được']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
schedule = [
    ['1', 'Tuần 1-2', 'Tìm hiểu yêu cầu, viết đề cương, thiết kế CSDL', 'Hoàn thành'],
    ['2', 'Tuần 3-5', 'Xây dựng giao diện và chức năng chính', 'Hoàn thành'],
    ['3', 'Tuần 6-8', 'Tích hợp Speaking AI, Wallet, Membership', 'Hoàn thành'],
    ['4', 'Tuần 9-10', 'Kiểm thử, sửa lỗi, viết UML + wireframe', 'Hoàn thành'],
    ['5', 'Tuần 11-12', 'Hoàn thiện báo cáo, chuẩn bị bảo vệ', 'Hoàn thành'],
]
for r, data in enumerate(schedule):
    for c, text in enumerate(data):
        cell = table.rows[r+1].cells[c]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

page_break()

# ============================================================
# CHƯƠNG 1: TỔNG QUAN VỀ ĐỀ TÀI
# ============================================================
heading1("TỔNG QUAN VỀ ĐỀ TÀI")

level1("Giới thiệu đề tài")
body_text("Trong bối cảnh toàn cầu hóa và hội nhập quốc tế ngày càng sâu rộng, tiếng Anh đã trở thành ngôn ngữ giao tiếp quốc tế không thể thiếu. Tại Việt Nam, nhu cầu học tiếng Anh ngày càng tăng cao, đặc biệt ở giới trẻ và sinh viên. Tuy nhiên, việc tiếp cận các khóa học tiếng Anh chất lượng với chi phí hợp lý vẫn còn là thách thức đối với nhiều người.")
body_text("Sự phát triển nhanh chóng của công nghệ thông tin và Internet đã mở ra cơ hội xây dựng các nền tảng học tập trực tuyến (e-learning), cho phép người học tiếp cận kiến thức mọi lúc, mọi nơi. Các ứng dụng học tiếng Anh trực tuyến giúp người dùng chủ động trong việc học, tiết kiệm thời gian và chi phí so với phương pháp truyền thống.")
body_text("Đề tài \"Xây dựng website học tiếng Anh trực tuyến English Learning\" được thực hiện nhằm phát triển một nền tảng web toàn diện, hỗ trợ người dùng học tiếng Anh qua nhiều kỹ năng: từ vựng, ngữ pháp, nghe, đọc, nói. Đặc biệt, hệ thống tích hợp trí tuệ nhân tạo (AI) để đánh giá kỹ năng nói, tạo trải nghiệm học tập tương tác và hiệu quả hơn.")

level1("Lý do chọn đề tài")
body_text("Việc chọn đề tài này xuất phát từ những lý do sau:")
bullet("Nhu cầu thực tế: Thị trường học tiếng Anh online tại Việt Nam đang tăng trưởng mạnh mẽ. Các nền tảng như Duolingo, ELSA Speak đã chứng minh tính khả thi và hiệu quả của việc học ngoại ngữ qua ứng dụng công nghệ.")
bullet("Ứng dụng công nghệ AI: Sự phát triển của AI (đặc biệt là GPT và Web Speech API) cho phép đánh giá kỹ năng nói tự động, mở ra hướng tiếp cận mới trong giáo dục ngoại ngữ.")
bullet("Tích hợp đa kỹ năng: Khác với các ứng dụng chỉ tập trung một kỹ năng, English Learning tích hợp tất cả 4 kỹ năng (Nghe, Nói, Đọc, Viết) và thêm từ vựng + ngữ pháp trong một nền tảng duy nhất.")
bullet("Hệ thống quản lý tài chính: Tích hợp ví điện tử, nạp tiền qua QR code, mua gói membership – cung cấp mô hình kinh doanh bền vững cho nền tảng giáo dục.")
bullet("Phù hợp với chuyên ngành: Đề tài cho phép áp dụng toàn diện các kiến thức đã học: lập trình web (PHP, HTML, CSS, JS), cơ sở dữ liệu (MySQL), mô hình MVC, và tích hợp API.")

level1("Mục tiêu đề tài")
level2("Mục tiêu tổng quát")
body_text("Xây dựng một website học tiếng Anh trực tuyến hoàn chỉnh, có giao diện thân thiện, nhiều tính năng học tập tương tác, và hệ thống quản trị + quản lý tài chính chuyên nghiệp.")

level2("Mục tiêu cụ thể")
bullet("Xây dựng hệ thống quản lý nội dung học tập: chủ đề từ vựng, bài học, bài kiểm tra, ngữ pháp.")
bullet("Phát triển tính năng Flashcard giúp người dùng học từ vựng hiệu quả.")
bullet("Tích hợp AI đánh giá kỹ năng nói (Speaking) với điểm số chi tiết: Pronunciation, Fluency, Accuracy.")
bullet("Xây dựng hệ thống bài kiểm tra đa dạng: Quiz, Listening, Reading với đếm thời gian.")
bullet("Phát triển hệ thống ví điện tử: nạp tiền, mua gói Pro, rút tiền, hoàn tiền.")
bullet("Xây dựng hệ thống membership (Free/Pro) với mã kích hoạt và thanh toán qua ví.")
bullet("Phát triển dashboard thống kê tiến độ học tập, bảng xếp hạng, hệ thống XP + Level.")
bullet("Xây dựng trang quản trị (Admin) với đầy đủ chức năng CRUD và quản lý tài chính.")

level1("Phạm vi đề tài")
body_text("Đề tài tập trung vào việc xây dựng và phát triển ứng dụng web với các phạm vi sau:")
bullet("Nền tảng: Website chạy trên trình duyệt (responsive, tương thích PC và mobile).")
bullet("Ngôn ngữ và công nghệ: PHP (thuần MVC), MySQL, HTML5, CSS3, JavaScript, AJAX.")
bullet("Tích hợp bên thứ ba: Google OAuth 2.0, OpenAI GPT API, Web Speech API, VietQR.")
bullet("Đối tượng sử dụng: Học sinh, sinh viên, người đi làm muốn học tiếng Anh trực tuyến.")
bullet("Phạm vi chức năng: Học từ vựng, ngữ pháp, luyện nói, làm bài kiểm tra, quản lý ví, mua gói membership.")
page_break()

# ============================================================
# CHƯƠNG 2: LÝ THUYẾT ÁP DỤNG
# ============================================================
heading1("LÝ THUYẾT ÁP DỤNG THEO LĨNH VỰC THỰC TẬP")

level1("Tổng quan về đề tài")

level2("Tổng quan về E-Learning")
body_text("E-Learning (Electronic Learning) là hình thức đào tạo trực tuyến sử dụng công nghệ thông tin và Internet để truyền tải nội dung học tập. E-Learning cho phép người học tiếp cận bài giảng, tài liệu, bài kiểm tra mọi lúc mọi nơi thông qua các thiết bị kết nối Internet.")
body_text("Ưu điểm của E-Learning bao gồm: linh hoạt về thời gian và địa điểm học, chi phí thấp hơn so với đào tạo truyền thống, khả năng cá nhân hóa lộ trình học tập, và tích hợp công nghệ AI để nâng cao trải nghiệm.")
body_text("Tại Việt Nam, thị trường E-Learning đang phát triển mạnh với các nền tảng như VioEdu, Topica, Moon.vn. Riêng lĩnh vực học tiếng Anh online, các ứng dụng như Duolingo, ELSA Speak, Cake đã thu hút hàng triệu người dùng Việt Nam.")

level2("Tổng quan về ứng dụng học tiếng Anh trực tuyến")
body_text("Các ứng dụng học tiếng Anh trực tuyến hiện nay thường tập trung vào các phương pháp sau:")
bullet("Gamification: Áp dụng các yếu tố trò chơi (điểm XP, streak, level, leaderboard) để tạo động lực học tập.")
bullet("Spaced Repetition: Lặp lại ngắt quãng – ôn tập từ vựng theo chu kỳ tối ưu để ghi nhớ lâu dài.")
bullet("AI-powered Assessment: Sử dụng AI để đánh giá phát âm, ngữ pháp, và cung cấp phản hồi tức thì.")
bullet("Multi-skill Integration: Tích hợp đa kỹ năng (Nghe, Nói, Đọc, Viết) trong một nền tảng duy nhất.")
body_text("English Learning áp dụng tất cả các phương pháp trên, kết hợp với hệ thống quản lý membership và ví điện tử, tạo ra một nền tảng học tập toàn diện.")

level1("Cơ sở lý thuyết")

level2("Ngôn ngữ lập trình PHP")
body_text("PHP (PHP: Hypertext Preprocessor) là ngôn ngữ lập trình kịch bản phía máy chủ (server-side scripting language) được sử dụng rộng rãi trong phát triển web. PHP được thiết kế để tạo ra các trang web động, xử lý form, quản lý session, cookie, và tương tác với cơ sở dữ liệu.")
body_text("Đặc điểm nổi bật của PHP:")
bullet("Mã nguồn mở, miễn phí, cộng đồng hỗ trợ lớn.")
bullet("Hỗ trợ nhiều hệ quản trị CSDL: MySQL, PostgreSQL, SQLite, Oracle.")
bullet("Tích hợp tốt với các web server: Apache, Nginx, IIS.")
bullet("Cú pháp đơn giản, dễ học, dễ triển khai.")
body_text("Trong đề tài này, PHP được sử dụng phiên bản 8.2 trở lên, áp dụng mô hình MVC để tổ chức mã nguồn một cách khoa học và dễ bảo trì.")

level2("Mô hình MVC (Model-View-Controller)")
body_text("MVC là một design pattern phổ biến trong phát triển web, chia ứng dụng thành 3 thành phần chính:")
bullet("Model: Xử lý logic nghiệp vụ và tương tác với cơ sở dữ liệu.")
bullet("View: Hiển thị giao diện người dùng (HTML, CSS, JavaScript).")
bullet("Controller: Tiếp nhận request từ người dùng, gọi Model xử lý, và chuyển kết quả đến View.")
body_text("Ưu điểm của MVC: tách biệt logic nghiệp vụ và giao diện, dễ bảo trì và mở rộng, hỗ trợ làm việc nhóm hiệu quả, tái sử dụng code.")

level2("Hệ quản trị CSDL MySQL")
body_text("MySQL là hệ quản trị cơ sở dữ liệu quan hệ (RDBMS) mã nguồn mở phổ biến nhất thế giới. MySQL sử dụng ngôn ngữ SQL để quản lý và truy vấn dữ liệu.")
body_text("Đặc điểm của MySQL:")
bullet("Hiệu suất cao, xử lý nhanh, hỗ trợ transaction (InnoDB engine).")
bullet("Hỗ trợ ACID đảm bảo tính toàn vẹn dữ liệu.")
bullet("Bảo mật tốt với hệ thống phân quyền chi tiết.")
bullet("Tích hợp tốt với PHP thông qua PDO (PHP Data Objects).")
body_text("Trong đề tài, MySQL được sử dụng với charset utf8mb4 để hỗ trợ đầy đủ ký tự tiếng Việt. Database gồm 23 bảng quản lý toàn bộ dữ liệu hệ thống.")

level2("HTML5, CSS3 và JavaScript")
body_text("HTML5 (HyperText Markup Language 5) là phiên bản mới nhất của ngôn ngữ đánh dấu dùng để cấu trúc nội dung trang web. HTML5 bổ sung các thẻ semantic (header, nav, section, article, footer) và hỗ trợ multimedia.")
body_text("CSS3 (Cascading Style Sheets 3) dùng để định dạng giao diện web. CSS3 bổ sung nhiều tính năng: Flexbox, Grid Layout, Animation, Gradient, Media Queries cho responsive design.")
body_text("JavaScript là ngôn ngữ lập trình phía client, cho phép tạo tương tác động trên trang web. Trong đề tài, JavaScript được sử dụng cho: AJAX, Web Speech API, Chart.js, và DOM manipulation.")

level2("Google OAuth 2.0")
body_text("OAuth 2.0 là giao thức ủy quyền cho phép ứng dụng bên thứ ba truy cập tài nguyên của người dùng mà không cần biết mật khẩu. Google OAuth 2.0 cho phép người dùng đăng nhập bằng tài khoản Google, đơn giản hóa quy trình đăng ký/đăng nhập và tăng bảo mật.")

level2("OpenAI GPT API")
body_text("OpenAI GPT (Generative Pre-trained Transformer) là mô hình ngôn ngữ lớn (LLM) có khả năng hiểu và tạo văn bản tự nhiên. Trong đề tài, GPT API được sử dụng để đánh giá kỹ năng nói: phân tích transcript từ Web Speech API, chấm điểm Pronunciation, Fluency, Accuracy và cung cấp feedback chi tiết.")

level2("Web Speech API")
body_text("Web Speech API là API trình duyệt cho phép tích hợp tính năng nhận dạng giọng nói (Speech Recognition) và tổng hợp giọng nói (Speech Synthesis) trực tiếp trên trình duyệt web. Trong đề tài, Speech Recognition API được sử dụng để ghi nhận giọng nói của người dùng khi luyện nói tiếng Anh, chuyển đổi thành văn bản rồi gửi lên server để AI đánh giá.")

level2("VietQR – Thanh toán QR Code")
body_text("VietQR là tiêu chuẩn QR Code thanh toán được Ngân hàng Nhà nước Việt Nam ban hành. VietQR API cho phép tạo mã QR chuyển khoản tự động với đầy đủ thông tin: số tài khoản, ngân hàng, số tiền, nội dung chuyển khoản. Trong đề tài, VietQR được tích hợp vào tính năng nạp tiền vào ví.")
page_break()

# ============================================================
# CHƯƠNG 3: NỘI DUNG THỰC TẬP
# ============================================================
heading1("NỘI DUNG THỰC TẬP")

level1("Phân tích và thiết kế hệ thống")

level2("Sơ đồ Use Case")
body_text("Sơ đồ Use Case mô tả các chức năng chính của hệ thống và mối quan hệ giữa các actor với hệ thống. Hệ thống English Learning có 4 actor chính:")
bullet("Guest (Khách): Người dùng chưa đăng nhập – xem chủ đề, đăng ký, đăng nhập.")
bullet("User (Học viên Free): Đã đăng nhập – học từ vựng, quiz, flashcard, ngữ pháp, bookmark, quản lý ví.")
bullet("Pro User (Hội viên Pro): Kế thừa quyền User + Listening, Reading, Speaking.")
bullet("Admin (Quản trị viên): Kế thừa quyền Pro + quản lý hệ thống, duyệt đơn, quản lý ví.")

add_image(os.path.join(UML_DIR, "use_case.png"), "Hình 3.1: Sơ đồ Use Case tổng quan", Inches(5.0))

level2("Sơ đồ Sequence")
body_text("Sơ đồ Sequence mô tả luồng tương tác giữa các thành phần trong hệ thống theo thời gian.")

level3("a) Luồng đăng nhập")
add_image(os.path.join(UML_DIR, "sequence_login.png"), "Hình 3.2: Sequence Diagram – Đăng nhập", Inches(5.5))

level3("b) Luồng nạp tiền vào ví")
add_image(os.path.join(UML_DIR, "sequence_deposit.png"), "Hình 3.3: Sequence Diagram – Nạp tiền vào ví", Inches(5.5))

level3("c) Luồng mua gói Pro bằng ví")
add_image(os.path.join(UML_DIR, "sequence_purchase.png"), "Hình 3.4: Sequence Diagram – Mua gói Pro", Inches(5.5))

level3("d) Luồng hủy đơn và hoàn tiền")
add_image(os.path.join(UML_DIR, "sequence_cancel.png"), "Hình 3.5: Sequence Diagram – Hủy đơn + Hoàn tiền", Inches(5.5))

level2("Sơ đồ State (Trạng thái)")
body_text("Sơ đồ State mô tả các trạng thái có thể có của một đối tượng và sự chuyển đổi giữa chúng.")

level3("a) Trạng thái Đơn hàng")
add_image(os.path.join(UML_DIR, "state_order.png"), "Hình 3.6: State Diagram – Membership Order", Inches(4.0))

level3("b) Trạng thái Giao dịch Ví")
add_image(os.path.join(UML_DIR, "state_wallet.png"), "Hình 3.7: State Diagram – Wallet Transaction", Inches(4.0))

level3("c) Trạng thái Support Ticket")
add_image(os.path.join(UML_DIR, "state_ticket.png"), "Hình 3.8: State Diagram – Support Ticket", Inches(4.0))

level2("Thiết kế cơ sở dữ liệu")
body_text("Hệ thống sử dụng MySQL với 23 bảng chính. Dưới đây là danh sách các bảng và mô tả:")

db_table = doc.add_table(rows=16, cols=3)
db_table.style = 'Table Grid'
db_headers = ['STT', 'Tên bảng', 'Mô tả']
for i, h in enumerate(db_headers):
    cell = db_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
db_data = [
    ['1', 'users', 'Thông tin người dùng (username, email, balance, membership, XP...)'],
    ['2', 'topics', 'Chủ đề học (Travel, Food, Business...)'],
    ['3', 'vocabularies', 'Từ vựng (word, pronunciation, meaning, example)'],
    ['4', 'lessons / lesson_contents', 'Bài học + nội dung đa phương tiện'],
    ['5', 'tests / questions', 'Bài kiểm tra + câu hỏi (quiz, listening, reading)'],
    ['6', 'grammar_lessons / grammar_questions', 'Bài ngữ pháp + câu hỏi quiz'],
    ['7', 'speaking_prompts / speaking_attempts', 'Đề luyện nói + kết quả AI chấm điểm'],
    ['8', 'membership_plans', 'Gói Pro (1 tháng, 3 tháng, 12 tháng)'],
    ['9', 'membership_orders', 'Đơn mua gói (pending/completed/cancelled)'],
    ['10', 'wallet_transactions', 'Giao dịch ví (deposit/purchase/refund/withdraw)'],
    ['11', 'support_tickets', 'Ticket hỗ trợ từ người dùng'],
    ['12', 'activation_codes', 'Mã kích hoạt Pro'],
    ['13', 'bookmarks', 'Từ vựng đã bookmark'],
    ['14', 'user_progress', 'Tiến độ học theo chủ đề'],
    ['15', 'xp_history', 'Lịch sử XP + hoạt động'],
]
for r, data in enumerate(db_data):
    for c, text in enumerate(data):
        cell = db_table.rows[r+1].cells[c]
        cell.text = text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

level1("Xây dựng giao diện ứng dụng")
body_text("Giao diện website English Learning được thiết kế theo phong cách hiện đại, sử dụng tone màu xanh dương làm chủ đạo, phối hợp với nền sáng tạo cảm giác thân thiện và dễ nhìn. Giao diện responsive, tương thích cả PC và mobile.")

screenshots = [
    ("Trang đăng nhập", "01_login.png", "Hình 3.9",
     "Cho phép người dùng đăng nhập bằng username/password hoặc tài khoản Google."),
    ("Trang chủ", "02_homepage_top.png", "Hình 3.10",
     "Hiển thị banner giới thiệu, thống kê nhanh và danh sách chủ đề học nổi bật."),
    ("Danh sách chủ đề", "04_topics.png", "Hình 3.11",
     "Grid các chủ đề theo level: Beginner, Intermediate, Advanced."),
    ("Chi tiết chủ đề", "05_topic_detail.png", "Hình 3.12",
     "Danh sách từ vựng, nút học Flashcard, làm bài test, danh sách bài học."),
    ("Flashcard", "07_flashcard.png", "Hình 3.13",
     "Thẻ lật hiển thị từ vựng: mặt trước từ tiếng Anh + IPA, mặt sau nghĩa + ví dụ."),
    ("Bài kiểm tra", "09_test.png", "Hình 3.14",
     "Danh sách bài test: Quiz (Free), Listening + Reading (Pro), có đếm thời gian."),
    ("Luyện nói (Speaking)", "10_speaking.png", "Hình 3.15",
     "Ghi âm giọng nói, AI chấm điểm Pronunciation, Fluency, Accuracy + feedback."),
    ("Dashboard", "11_dashboard.png", "Hình 3.16",
     "Thống kê tiến độ: streak, XP, Level, mục tiêu hàng ngày, biểu đồ tiến độ."),
    ("Ví điện tử", "15_wallet.png", "Hình 3.17",
     "Số dư, nút nạp/rút/mua gói, lịch sử giao dịch chi tiết."),
    ("Nạp tiền", "16_deposit.png", "Hình 3.18",
     "Tích hợp VietQR: nhập số tiền, hệ thống tự tạo QR code chuyển khoản."),
    ("Nâng cấp Pro", "18_membership_pricing.png", "Hình 3.19",
     "Bảng giá 3 gói, so sánh Free vs Pro, thanh toán bằng ví."),
]

for title, img_file, fig_num, desc in screenshots:
    level2(title)
    body_text(desc)
    add_image(os.path.join(SCREENSHOTS_DIR, img_file), f"{fig_num}: Giao diện {title}", Inches(5.2))

level1("Giao diện quản trị (Admin)")
body_text("Trang quản trị dành cho Admin với đầy đủ chức năng quản lý hệ thống:")

admin_screenshots = [
    ("Dashboard Admin", "22_admin_dashboard.png", "Hình 3.20",
     "Thống kê tổng quan: Users, Pro Members, Chủ đề, Lượt làm bài, Tickets. Kèm 4 biểu đồ."),
    ("Quản lý Đơn nâng cấp", "27_admin_orders.png", "Hình 3.21",
     "Danh sách đơn hàng membership. Admin duyệt hoặc từ chối đơn pending."),
    ("Quản lý Ví", "29_admin_wallet.png", "Hình 3.22",
     "Danh sách giao dịch ví. Admin duyệt/từ chối nạp/rút tiền."),
    ("Quản lý Tickets", "28_admin_tickets.png", "Hình 3.23",
     "Danh sách tickets. Admin phản hồi, duyệt hủy đơn + hoàn tiền tự động."),
]

for title, img_file, fig_num, desc in admin_screenshots:
    level2(title)
    body_text(desc)
    add_image(os.path.join(SCREENSHOTS_DIR, img_file), f"{fig_num}: Giao diện {title}", Inches(5.2))

level1("Xây dựng chức năng")
body_text("Hệ thống English Learning bao gồm các nhóm chức năng chính sau:")

level2("Quản lý tài khoản")
bullet("Đăng ký tài khoản mới (username, email, password).")
bullet("Đăng nhập bằng username/password hoặc Google OAuth 2.0.")
bullet("Quản lý hồ sơ: cập nhật thông tin, đổi avatar, đổi mật khẩu.")
bullet("Phân quyền: Student (Free/Pro) và Admin.")

level2("Hệ thống học tập")
bullet("Chủ đề từ vựng: 6 chủ đề, mỗi chủ đề có từ vựng, bài học, bài kiểm tra.")
bullet("Flashcard: Thẻ lật hiển thị từ vựng, bookmark từ yêu thích.")
bullet("Bài học (Lesson): Nội dung đa phương tiện, đánh giá sao.")
bullet("Ngữ pháp: 10 bài ngữ pháp với quiz trắc nghiệm.")
bullet("Bài kiểm tra: Quiz (Free), Listening + Reading (Pro), đếm thời gian, chấm điểm tự động.")
bullet("Luyện nói: Ghi âm → Speech-to-Text → AI đánh giá → Feedback chi tiết.")

level2("Hệ thống Gamification")
bullet("XP (Experience Points): Nhận XP khi hoàn thành hoạt động học tập.")
bullet("Level: Nâng level khi đạt đủ XP (mỗi level cần 100 XP).")
bullet("Streak: Đếm số ngày học liên tiếp, tặng bonus XP.")
bullet("Leaderboard: Bảng xếp hạng top users theo XP.")
bullet("Daily Goal: Mục tiêu XP hàng ngày (mặc định 50 XP).")

level2("Hệ thống tài chính")
bullet("Ví điện tử: Mỗi user có balance riêng, ghi nhận tất cả giao dịch.")
bullet("Nạp tiền: User chuyển khoản → Admin duyệt → Cộng balance.")
bullet("Mua gói Pro: Trừ balance → Tạo order completed → Nâng cấp Pro.")
bullet("Rút tiền: User yêu cầu → Admin chuyển khoản → Trừ balance.")
bullet("Hoàn tiền: Hủy đơn → Admin duyệt → Hoàn tiền vào ví (tỷ lệ theo thời gian).")
bullet("Bảo mật: Sử dụng transaction + row locking (FOR UPDATE) tránh race condition.")

level2("Hệ thống hỗ trợ")
bullet("Ticket hỗ trợ: User gửi ticket (Hỗ trợ chung, Hủy đơn, Báo lỗi, Góp ý).")
bullet("Admin phản hồi: Trả lời ticket, đổi trạng thái (Open → Resolved → Closed).")
bullet("Hủy đơn + Hoàn tiền: Chính sách theo thời gian (24h: 100%, 1-7 ngày: 50%).")

level1("Kiểm thử")
body_text("Quá trình kiểm thử được thực hiện trên 30 trang web, bao gồm cả giao diện người dùng và quản trị. Kết quả kiểm thử:")
bullet("30/30 trang hoạt động bình thường, không lỗi PHP.")
bullet("Vietnamese encoding hiển thị chính xác trên tất cả các trang.")
bullet("Wallet + Membership flows hoạt động đúng: nạp tiền, mua gói, hủy đơn, hoàn tiền.")
bullet("Google OAuth đăng nhập thành công.")
bullet("Speaking AI chấm điểm chính xác.")
bullet("Responsive design hoạt động tốt trên PC và mobile.")
page_break()

# ============================================================
# CHƯƠNG 4: KẾT LUẬN VÀ KIẾN NGHỊ
# ============================================================
heading1("KẾT LUẬN VÀ KIẾN NGHỊ")

level1("Kết quả đạt được")
level2("Ưu điểm")
bullet("Hoàn thành đầy đủ các chức năng đề ra: học từ vựng, ngữ pháp, bài kiểm tra, luyện nói, flashcard, bookmark.")
bullet("Tích hợp thành công AI (OpenAI GPT) vào đánh giá kỹ năng nói – điểm nổi bật của đề tài.")
bullet("Hệ thống ví điện tử + membership chuyên nghiệp với đầy đủ flow: nạp, mua, rút, hoàn tiền.")
bullet("Giao diện người dùng đẹp, hiện đại, responsive, trải nghiệm mượt mà.")
bullet("Hệ thống gamification (XP, Level, Streak, Leaderboard) tạo động lực học tập.")
bullet("Bảo mật giao dịch tài chính: sử dụng database transaction + row locking.")
bullet("Code sạch, tổ chức theo mô hình MVC, dễ bảo trì và mở rộng.")
bullet("Đăng nhập bằng Google OAuth 2.0 tiện lợi cho người dùng.")

level2("Hạn chế")
bullet("Tính năng nạp tiền chưa tích hợp tự động (cần Admin duyệt thủ công).")
bullet("Chưa có ứng dụng mobile native (chỉ hoạt động trên trình duyệt web).")
bullet("Nội dung học tập còn hạn chế (6 chủ đề), cần bổ sung thêm.")
bullet("Chưa có tính năng chat/forum cho người dùng trao đổi.")
bullet("Speaking chỉ hoạt động trên trình duyệt hỗ trợ Web Speech API (Chrome, Edge).")

level1("Hạn chế")
body_text("Do thời gian và kiến thức còn hạn chế, đề tài chưa thể triển khai trên môi trường hosting thật và chưa tích hợp cổng thanh toán tự động. Nội dung học tập cần được bổ sung thêm nhiều chủ đề hơn để phục vụ đa dạng nhu cầu người học.")

level1("Kiến nghị (Hướng phát triển trong tương lai)")
bullet("Tích hợp cổng thanh toán tự động: VNPay, Momo, ZaloPay để xác nhận nạp tiền real-time.")
bullet("Phát triển ứng dụng mobile bằng React Native hoặc Flutter.")
bullet("Bổ sung thêm nội dung: nhiều chủ đề hơn, video bài giảng, podcast tiếng Anh.")
bullet("Tích hợp Spaced Repetition System (SRS) để tối ưu hóa việc ôn tập từ vựng.")
bullet("Thêm tính năng học nhóm, diễn đàn trao đổi giữa người dùng.")
bullet("Sử dụng AI nâng cao: chatbot hội thoại, tự động tạo đề thi.")
bullet("Deploy lên hosting thật (VPS, Cloud) để phục vụ người dùng thực tế.")
page_break()

# ============================================================
# TÀI LIỆU THAM KHẢO
# ============================================================
heading1("TÀI LIỆU THAM KHẢO")

body_text("Tài liệu tham khảo xếp theo thứ tự ABC họ tên tác giả:")

normal("Tiếng Việt", True)
normal("[1]\tNguyễn Văn Hiệp (2023), Giáo trình Lập trình Web với PHP và MySQL, NXB Giáo dục Việt Nam.")

normal("Tiếng Anh", True)
normal("[2]\tMatt Zandstra (2021), PHP 8 Objects, Patterns, and Practice, Apress, 6th Edition.")
normal("[3]\tLuke Welling, Laura Thomson (2021), PHP and MySQL Web Development, Addison-Wesley, 5th Edition.")
normal("[4]\tRobin Nixon (2021), Learning PHP, MySQL & JavaScript, O'Reilly Media, 6th Edition.")

normal("Trang web", True)
normal("[5]\tPHP Official Documentation. Địa chỉ: https://www.php.net/docs.php")
normal("[6]\tMySQL Reference Manual. Địa chỉ: https://dev.mysql.com/doc/refman/8.0/en/")
normal("[7]\tMDN Web Docs – HTML, CSS, JavaScript. Địa chỉ: https://developer.mozilla.org/")
normal("[8]\tGoogle OAuth 2.0 Documentation. Địa chỉ: https://developers.google.com/identity/protocols/oauth2")
normal("[9]\tOpenAI API Documentation. Địa chỉ: https://platform.openai.com/docs/")
normal("[10]\tWeb Speech API – MDN. Địa chỉ: https://developer.mozilla.org/en-US/docs/Web/API/Web_Speech_API")
normal("[11]\tVietQR API Documentation. Địa chỉ: https://www.vietqr.io/")
normal("[12]\tChart.js Documentation. Địa chỉ: https://www.chartjs.org/docs/")

# ============================================================
# SAVE
# ============================================================
doc.save(OUTPUT)
size = os.path.getsize(OUTPUT)
print(f"Done! Saved to: {OUTPUT}")
print(f"File size: {size/1024:.0f} KB ({size/1024/1024:.1f} MB)")
print(f"Paragraphs: {len(doc.paragraphs)}")
